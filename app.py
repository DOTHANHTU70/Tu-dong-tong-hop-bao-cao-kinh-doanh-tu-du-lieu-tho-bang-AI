from flask import Flask, render_template_string, request
import pandas as pd
import os
import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from datetime import datetime
import time
from docx import Document
from docx.shared import Inches, Pt

app = Flask(__name__)

# Thư mục lưu file upload & ảnh biểu đồ
UPLOAD_FOLDER = 'uploads'
STATIC_FOLDER = 'static'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(STATIC_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ====== 1) Hỗ trợ map tên cột linh hoạt ======
COLUMN_ALIASES = {
    "date": {"date", "ngay", "ngày", "order_date", "ngày bán", "time", "thoi_gian", "thời gian"},
    "product_id": {"product_id", "ma_sp", "mã_sp", "mã sản phẩm", "product", "item", "sku"},
    "revenue": {"revenue", "doanh_thu", "doanh thu", "sales", "amount", "gross_sales", "tiền bán"},
    "profit": {"profit", "loi_nhuan", "lợi nhuận", "net_profit", "lãi"}
}

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    cols = {c: c.strip().lower() for c in df.columns}
    df = df.rename(columns=cols)

    mapped = {}
    for std_col, aliases in COLUMN_ALIASES.items():
        for c in df.columns:
            if c in aliases:
                mapped[std_col] = c
                break

    required = {"date", "revenue", "profit", "product_id"}
    if not required.issubset(mapped.keys()):
        missing = sorted(list(required - set(mapped.keys())))
        raise ValueError(
            "File Excel thiếu cột bắt buộc hoặc tên cột không nhận diện được: "
            + ", ".join(missing)
            + ".\nGợi ý đặt tên cột ví dụ: date/ngày, product_id/mã sản phẩm, revenue/doanh thu, profit/lợi nhuận."
        )

    df = df[[mapped["date"], mapped["product_id"], mapped["revenue"], mapped["profit"]]].copy()
    df.columns = ["date", "product_id", "revenue", "profit"]
    return df

# ====== 2) Tiền xử lý + tính KPI ======
def preprocess_data_from_excel(file_path, sheet_name=0):
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        df = normalize_columns(df)

        for col in ["revenue", "profit"]:
            df[col] = (
                df[col]
                .apply(lambda x: str(x).replace(",", "").replace(" ", "") if pd.notna(x) else x)
            )
            df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

        df["date"] = pd.to_datetime(df["date"], errors="coerce")
        df = df.dropna(subset=["date"])
        df = df.fillna({"product_id": "UNKNOWN"})

        total_revenue = float(df["revenue"].sum())
        total_profit = float(df["profit"].sum())
        gross_margin = (total_profit / total_revenue) if total_revenue > 0 else 0.0

        monthly = (
            df.set_index("date")
              .sort_index()
              .resample("M")
              .agg({"revenue": "sum", "profit": "sum"})
        )
        avg_monthly_revenue = float(monthly["revenue"].mean()) if len(monthly) > 0 else 0.0

        mom_growth = None
        if len(monthly) >= 2:
            last = monthly["revenue"].iloc[-1]
            prev = monthly["revenue"].iloc[-2]
            mom_growth = ((last - prev) / prev) if prev != 0 else None

        top_products_series = (
            df.groupby("product_id")["revenue"].sum().sort_values(ascending=False).head(5)
        )
        top_selling_products = top_products_series.to_dict()

        return {
            "df": df,
            "monthly": monthly,
            "total_revenue": total_revenue,
            "total_profit": total_profit,
            "gross_margin": gross_margin,
            "average_monthly_revenue": avg_monthly_revenue,
            "mom_growth": mom_growth,
            "top_selling_products": top_selling_products,
        }
    except Exception as e:
        print(f"Lỗi khi xử lý dữ liệu: {e}")
        return None

# ====== 3) Vẽ biểu đồ và lưu ảnh ======
def generate_charts(processed):
    charts = {}
    if processed["monthly"] is not None and len(processed["monthly"]) > 0:
        fig1 = plt.figure()
        processed["monthly"]["revenue"].plot(marker="o")
        plt.title("Xu hướng doanh thu theo tháng")
        plt.xlabel("Tháng")
        plt.ylabel("Doanh thu")
        plt.tight_layout()
        chart1_path = os.path.join(STATIC_FOLDER, "revenue_trend.png")
        fig1.savefig(chart1_path)
        plt.close(fig1)
        charts["revenue_trend"] = "/static/revenue_trend.png"

    if processed["top_selling_products"]:
        s = pd.Series(processed["top_selling_products"])
        fig2 = plt.figure()
        s.sort_values(ascending=True).plot(kind="barh")
        plt.title("Top sản phẩm theo doanh thu")
        plt.xlabel("Doanh thu")
        plt.ylabel("Sản phẩm")
        plt.tight_layout()
        chart2_path = os.path.join(STATIC_FOLDER, "top_products.png")
        fig2.savefig(chart2_path)
        plt.close(fig2)
        charts["top_products"] = "/static/top_products.png"

    return charts

# ====== 4) “AI” sinh báo cáo tự động (rule-based) ======
def generate_report_with_ai_rule_based(data):
    total_rev = data["total_revenue"]
    total_profit = data["total_profit"]
    margin = data["gross_margin"]
    avg_rev = data["average_monthly_revenue"]
    mom = data["mom_growth"]
    top = data["top_selling_products"]

    lines = []
    lines.append("--- BÁO CÁO KINH DOANH TỰ ĐỘNG ---\n")
    lines.append("Kính gửi Ban lãnh đạo,\n")
    lines.append("Báo cáo dưới đây được tổng hợp tự động từ dữ liệu thô do phòng kinh doanh cung cấp.\n")
    lines.append("1) TỔNG QUAN KẾT QUẢ\n---------------------")
    lines.append(f"- Tổng doanh thu: {total_rev:,.0f} VNĐ")
    lines.append(f"- Tổng lợi nhuận: {total_profit:,.0f} VNĐ")
    lines.append(f"- Biên lợi nhuận gộp (ước tính): {margin*100:,.2f}%")
    lines.append(f"- Doanh thu trung bình/tháng: {avg_rev:,.0f} VNĐ")

    lines.append("\n2) XU HƯỚNG DOANH THU\n----------------------")
    if mom is None:
        lines.append("- Dữ liệu chưa đủ để tính tăng trưởng tháng gần nhất.")
    else:
        pct = mom * 100
        if pct > 0:
            lines.append(f"- Doanh thu tháng gần nhất tăng {pct:,.2f}% so với tháng liền trước.")
        elif pct < 0:
            lines.append(f"- Doanh thu tháng gần nhất giảm {abs(pct):,.2f}% so với tháng liền trước.")
        else:
            lines.append("- Doanh thu không đổi so với tháng trước.")

    lines.append("\n3) SẢN PHẨM CHỦ LỰC\n-------------------")
    if top:
        items = list(top.items())
        for i, (pid, rev) in enumerate(items[:3], start=1):
            lines.append(f"- Top {i}: {pid} — doanh thu {rev:,.0f} VNĐ")
    else:
        lines.append("- Chưa đủ dữ liệu sản phẩm để xếp hạng.")

    lines.append("\n4) KẾT LUẬN & ĐỀ XUẤT\n----------------------")
    if total_rev > 0 and margin >= 0.15:
        lines.append("- Biên lợi nhuận ở mức khỏe, có dư địa tăng trưởng.")
    elif total_rev > 0 and margin < 0.08:
        lines.append("- Biên lợi nhuận mỏng; cân nhắc tối ưu chi phí hoặc điều chỉnh giá bán.")
    else:
        lines.append("- Chưa đủ dữ liệu kết luận về hiệu quả tổng thể.")

    lines.append("- Đề xuất: tập trung vào nhóm sản phẩm Top, tối ưu kênh mang lại tăng trưởng, và theo dõi xu hướng MoM trong 2–3 tháng tới.")
    lines.append("\nTrân trọng.")
    return "\n".join(lines)

# ====== 4b) Xuất báo cáo DOCX ======
def export_report_docx(processed, charts, report_text, output_path):
    doc = Document()
    run = doc.add_paragraph().add_run("BÁO CÁO KINH DOANH TỰ ĐỘNG")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph(f"Thời điểm tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_heading("1) KPI Tổng quan", level=2)
    doc.add_paragraph(f"- Tổng doanh thu: {processed['total_revenue']:,.0f} VNĐ")
    doc.add_paragraph(f"- Tổng lợi nhuận: {processed['total_profit']:,.0f} VNĐ")
    doc.add_paragraph(f"- Biên lợi nhuận gộp: {processed['gross_margin']*100:,.2f}%")
    doc.add_paragraph(f"- Doanh thu trung bình/tháng: {processed['average_monthly_revenue']:,.0f} VNĐ")
    if processed["mom_growth"] is None:
        doc.add_paragraph("- Tăng trưởng MoM: Chưa đủ dữ liệu")
    else:
        doc.add_paragraph(f"- Tăng trưởng MoM: {processed['mom_growth']*100:,.2f}%")
    doc.add_heading("2) Top sản phẩm theo doanh thu", level=2)
    top = processed.get("top_selling_products", {})
    if top:
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Mã sản phẩm"
        hdr_cells[1].text = "Doanh thu (VNĐ)"
        for pid, rev in top.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(pid)
            row_cells[1].text = f"{rev:,.0f}"
    doc.add_heading("3) Biểu đồ", level=2)
    if charts.get("revenue_trend"):
        doc.add_paragraph("• Xu hướng doanh thu theo tháng")
        try:
            doc.add_picture(os.path.join(".", charts["revenue_trend"].lstrip("/")), width=Inches(5.5))
        except Exception:
            pass
    if charts.get("top_products"):
        doc.add_paragraph("• Top sản phẩm theo doanh thu")
        try:
            doc.add_picture(os.path.join(".", charts["top_products"].lstrip("/")), width=Inches(5.5))
        except Exception:
            pass
    doc.add_heading("4) Nhận định & đề xuất", level=2)
    for line in report_text.splitlines():
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path

# ====== 5) Giao diện ======
INDEX_HTML = '''
<!doctype html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<title>Ứng dụng Báo cáo AI</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<style>
    body { background: #f8f9fa; }
    .upload-area {
        border: 2px dashed #6c757d;
        border-radius: 12px;
        padding: 40px;
        text-align: center;
        background: white;
        transition: border 0.3s ease;
    }
    .upload-area.dragover {
        border-color: #0d6efd;
        background: #e9f2ff;
    }
    .upload-area input {
        display: none;
    }
    .upload-icon {
        font-size: 50px;
        color: #0d6efd;
    }
    .file-name {
        margin-top: 15px;
        font-weight: 500;
        color: #495057;
    }
</style>
</head>
<body>
<div class="container mt-5">
    <div class="card shadow-sm p-4">
        <h2 class="card-title text-center mb-4">Tải hoặc Kéo-Thả file Excel</h2>
        <p class="text-center text-muted">
            File Excel cần có các cột: ngày/date, doanh thu/revenue, lợi nhuận/profit, mã sản phẩm/product_id.
        </p>
        <form id="uploadForm" method="post" action="/upload" enctype="multipart/form-data">
            <div id="uploadArea" class="upload-area">
                <div class="upload-icon">📂</div>
                <p>Kéo và thả file Excel vào đây<br>hoặc bấm để chọn file</p>
                <input type="file" name="file" id="fileInput" accept=".xlsx,.csv" required>
                <p id="fileName" class="file-name"></p>
            </div>
            <div class="text-center mt-3">
                <button type="submit" class="btn btn-primary">Tạo Báo Cáo</button>
            </div>
        </form>
    </div>
</div>

<script>
    const uploadArea = document.getElementById("uploadArea");
    const fileInput = document.getElementById("fileInput");
    const fileNameDisplay = document.getElementById("fileName");

    // Bấm để mở file explorer
    uploadArea.addEventListener("click", () => fileInput.click());

    // Khi kéo file vào
    uploadArea.addEventListener("dragover", (e) => {
        e.preventDefault();
        uploadArea.classList.add("dragover");
    });

    uploadArea.addEventListener("dragleave", () => {
        uploadArea.classList.remove("dragover");
    });

    // Khi thả file
    uploadArea.addEventListener("drop", (e) => {
        e.preventDefault();
        uploadArea.classList.remove("dragover");
        if (e.dataTransfer.files.length > 0) {
            fileInput.files = e.dataTransfer.files;
            fileNameDisplay.textContent = "File đã chọn: " + fileInput.files[0].name;
        }
    });

    // Khi chọn file qua explorer
    fileInput.addEventListener("change", () => {
        if(fileInput.files.length > 0){
            fileNameDisplay.textContent = "File đã chọn: " + fileInput.files[0].name;
        }
    });
</script>
</body>
</html>
'''


RESULT_HTML = '''
<!doctype html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<title>Kết quả Báo cáo</title>
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light">
<div class="container mt-5">
    <div class="text-center mb-4">
        <h1>Kết quả Báo Cáo</h1>
        <a href="/" class="btn btn-secondary mt-2 me-2">Quay lại</a>
        {% if download_docx %}
          <a href="{{ download_docx }}" class="btn btn-success mt-2">Tải báo cáo (.docx)</a>
        {% endif %}
    </div>
    <div class="row g-4">
        <div class="col-md-6">
            <div class="card shadow-sm p-3">
                <h5>KPI Tổng quan</h5>
                <ul class="list-unstyled mt-2">
                    <li>Tổng doanh thu: {{ total_revenue | int }} VNĐ</li>
                    <li>Tổng lợi nhuận: {{ total_profit | int }} VNĐ</li>
                    <li>Biên lợi nhuận gộp: {{ (gross_margin*100)|round(2) }}%</li>
                    <li>Doanh thu trung bình/tháng: {{ average_monthly_revenue | int }} VNĐ</li>
                    {% if mom_growth is none %}
                        <li>Tăng trưởng MoM: <span class="text-danger">Chưa đủ dữ liệu</span></li>
                    {% else %}
                        <li>Tăng trưởng MoM: {{ (mom_growth*100)|round(2) }}%</li>
                    {% endif %}
                </ul>
            </div>
        </div>
        <div class="col-md-6">
            <div class="card shadow-sm p-3">
                <h5>Top sản phẩm theo doanh thu</h5>
                {% if top_selling_products %}
                <ol class="mt-2">
                    {% for pid, rev in top_selling_products.items() %}
                        <li>{{ pid }} — {{ rev | int }} VNĐ</li>
                    {% endfor %}
                </ol>
                {% else %}
                    <p class="text-danger">Chưa có dữ liệu sản phẩm.</p>
                {% endif %}
            </div>
        </div>
    </div>
    <div class="row mt-4 g-4">
        {% if charts.revenue_trend %}
        <div class="col-md-6">
            <div class="card shadow-sm p-3">
                <h6>Xu hướng doanh thu theo tháng</h6>
                <img src="{{ charts.revenue_trend }}" class="img-fluid mt-2">
            </div>
        </div>
        {% endif %}
        {% if charts.top_products %}
        <div class="col-md-6">
            <div class="card shadow-sm p-3">
                <h6>Top sản phẩm</h6>
                <img src="{{ charts.top_products }}" class="img-fluid mt-2">
            </div>
        </div>
        {% endif %}
    </div>
    <div class="card shadow-sm p-3 mt-4">
        <h5>Báo cáo chi tiết</h5>
        <pre class="mt-2">{{ report_content }}</pre>
    </div>
</div>
</body>
</html>
'''

# ====== 6) Routes ======
@app.route('/')
def index():
    return render_template_string(INDEX_HTML)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "Không tìm thấy file!"
    file = request.files['file']
    if file.filename == '':
        return "Không có file nào được chọn."

    ext = os.path.splitext(file.filename)[1].lower()
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    # ===== Xử lý dữ liệu theo loại file =====
    if ext == ".xlsx":
        processed = preprocess_data_from_excel(file_path)
    elif ext == ".csv":
        try:
            df = pd.read_csv(file_path)

            # Chuẩn hóa để tận dụng lại hàm preprocess
            tmp_path = file_path + ".xlsx"
            df.to_excel(tmp_path, index=False)
            processed = preprocess_data_from_excel(tmp_path)

            # Xóa file tạm
            os.remove(tmp_path)
        except Exception as e:
            return f"Lỗi khi đọc file CSV: {e}"
    else:
        return "⚠️ Hệ thống chỉ hỗ trợ file .xlsx hoặc .csv"

    # Xóa file gốc sau khi xử lý xong
    try:
        os.remove(file_path)
    except Exception:
        pass

    if not processed:
        return "Đã xảy ra lỗi khi xử lý dữ liệu. Vui lòng kiểm tra lại file của bạn."

    # Vẽ biểu đồ
    charts = generate_charts(processed)

    # Báo cáo “AI”
    report_content = generate_report_with_ai_rule_based(processed)

    # ===== Xuất báo cáo DOCX =====
    ts = time.strftime("%Y%m%d_%H%M%S")
    docx_filename = f"bao_cao_kinh_doanh_{ts}.docx"
    docx_path = os.path.join(STATIC_FOLDER, docx_filename)
    export_report_docx(processed, charts, report_content, docx_path)
    download_docx_url = f"/static/{docx_filename}"

    # Render ra giao diện
    return render_template_string(
        RESULT_HTML,
        total_revenue=processed["total_revenue"],
        total_profit=processed["total_profit"],
        gross_margin=processed["gross_margin"],
        average_monthly_revenue=processed["average_monthly_revenue"],
        mom_growth=processed["mom_growth"],
        top_selling_products=processed["top_selling_products"],
        charts=charts,
        report_content=report_content,
        download_docx=download_docx_url
    )

if __name__ == '__main__':
    app.run(debug=True)
